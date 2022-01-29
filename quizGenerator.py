''' quizGenerator.py - To generate quizes according to the vocabulary table

The programme first get the vocabulary table from a xlsx document,
and then generate the correspoding dictation quizes.

Typical using examples:

quizGenerator.work()

'''
import xlrd
import re
import os
import glob
import docx
import random
import currentTime
import time
import sys

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

currentDate = time.strftime("%Y-%m-%d", time.localtime())

# regex family
wordRegex = re.compile(r'^[a-zA-Z]+$')  # English words
chineseRegex = re.compile(r'[\u4e00-\u9fa5]+')  # Chinese characters
rangeRegex = re.compile(r'[0-9]+-[0-9]+') # a range e.g. 1-100
numRegex = re.compile(r'[0-9]+')    # pure integers

# the main funciton(find the xlsx source)
def work(openPath, binPath, versionCnt, wordRangeBegin, wordRangeEnd, mode, questionCnt):
    file = openPath
    #   give the user an idea of which file is being worked on, and a chance to quit
    bsName = os.path.basename(file)
    sourceName = bsName.split('.')[0]
    try:
        gen(file, sourceName, binPath, versionCnt, wordRangeBegin, wordRangeEnd, mode, questionCnt)
    except Exception as exc:
        print('This workbook has failed: ' + file + '\n\n' + str(exc))
        

def gen(fname, sourceName, binPath, versionCnt, wordRangeBegin, wordRangeEnd, mode, questionCnt):
    '''get the words from the xlsx document which is directed by fname
    and then goto quizPrint() to print
    gen() ask a lot of personalizing questions

    '''
    wb = xlrd.open_workbook(fname)
    ws = wb.sheet_by_index(0)
    
    # get the available row and column numbers
    rowNum = ws.nrows
    colNum = ws.ncols
    wordCol = -1
    
    # find where the English words are
    for i in range(colNum):
        for j in range(1,rowNum):
            val = str(ws.col_values(i)[j])
            val = ''.join(val.split())
            if wordRegex.search(val):
                wordCol = i
                break
        if wordCol != -1:
            break
    if wordCol == -1:
        raise Exception('Word no find.')

    # find where the translations are
    transCol = -1
    for i in range(colNum):
        if i == wordCol:
            continue
        for j in range(1,rowNum):
            val = str(ws.col_values(i)[j])
            val = ''.join(val.split())
            if chineseRegex.search(val):
                transCol = i
                #print('DEBUG:' + val)
                break
        if transCol != -1:
            break
    if transCol == -1:
        raise Exception('Translation no find.')

    
    # generate the quiz form

    # get the number of students
    stuNum = versionCnt
    # get the range of words， and put them to head and tail
    head = wordRangeBegin
    tail = wordRangeEnd
    # get the user need eng to chn or chn to eng
    opt = mode
    questionNum = questionCnt
    
        
    # do a loop, generate a different quiz for every student
    for i in range(stuNum):
        # read in all the words and its translations
        wordList = {}
        sampleList = random.sample(range(head, tail + 1), questionNum)
        for wordID in sampleList:   # loop all the rows, search for the word
            key = str(ws.col_values(wordCol)[wordID])
            key = ' '.join(key.split())
            val = str(ws.col_values(transCol)[wordID])
            val = ' '.join(val.split())
            wordList[key] = val  # use a dictionary to storage: key for the word, value for the translation
        printQuiz(i, wordList, opt, head, tail, binPath, sourceName, answerOn=False)
        printQuiz(i, wordList, opt, head, tail, binPath, sourceName,answerOn=True)

  
def printQuiz(stu, wordList, opt, head, tail, binPath, sourceName, answerOn=False):
    '''print the quiz beautifully into docx files
    '''
    # create a document
    doc = docx.Document()
    
    # change the style
    doc.styles['Normal'].font.name = 'Cambria'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(15)
    doc.styles['Normal'].font.color.rgb = RGBColor(0,0,0)
    
    # create a style
    style_song = doc.styles.add_style('Song', WD_STYLE_TYPE.CHARACTER)
    style_song.font.name = '微软雅黑'
    doc.styles['Song']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    
    # add the headings
    h1 = doc.add_heading('', level=0)
    h1.add_run('默写纸', style='Song')
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    h2 = doc.add_heading('', level=4)
    h2.add_run('日期：' + currentDate, style='Song')
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    
    # divide into two sections
    section = doc.sections[0]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'2')

    # output
    
    #   设置行距和字体，后面附上下划线。
    #   TODO:如果超过一页纸要进行警告
    #   check for the opt to decide which mode will be executed
    
    for i in wordList:
        para = doc.add_paragraph(style = 'List Number')
        para.paragraph_format.line_spacing = Pt(30)

        if opt == 0:    # 英翻中
            st = i
            ans = wordList[i]
        else:   # 中翻英
            st = wordList[i]
            ans = i

        # check if this is an answer sheet
        if answerOn == False :
            ans = '\t\t\t'
        else:
            st += ':'   #   pretty print

        # add run in the paragraph    
        run1 = para.add_run()
        run1.text = st
        run1.font.size = Pt(15)
        
        run2 = para.add_run()
        run2.text = ans
        run2.font.size = Pt(15)
        run2.font.underline = True
           
    # save
    # 把xlsx的文件名写在生成的默写纸里面（默写纸_高频词汇表_1-100词_版本1_YYYY-MM-DD hh.mm.xlsx）
    
    #   show the range of words in the file name
    rangeName = str(head) + '-' + str(tail) + '词'
    #   show the version in the file name
    versionName = '版本' + str(stu + 1)
    #   get the current time
    cT = currentTime.getTime()
    #   check whether its a quiz or an answerbook
    if answerOn == False:
        '''默写纸_高频词汇表_1-100词_版本1_YYYY-MM-DD hh.mm.xlsx
        '''
        path = '_'.join(['默写纸', sourceName, rangeName, versionName, cT]) + '.docx'
    else:
        '''答案_高频词汇表_1-100词_版本1_YYYY-MM-DD hh.mm.xlsx
        '''
        path = '_'.join(['答案', sourceName, rangeName, versionName, cT]) + '.docx'
    #   make sure the directory exists
    if not os.path.exists(binPath):
        os.makedirs(binPath)
    #   save
    if not (binPath.endswith("\\") or binPath.endswith("/")):
        binPath += "\\"
    path = binPath + path
    doc.save(path)
    print(os.path.abspath(path))
    

if __name__=='__main__':
    if len(sys.argv) < 7:
        sys.exit(-1)
    openPath = sys.argv[1]
    binPath = sys.argv[2]
    versionCnt = int(sys.argv[3])
    wordRangeBegin = int(sys.argv[4])
    wordRangeEnd = int(sys.argv[5])
    mode = int(sys.argv[6])
    questionCnt = int(sys.argv[7])
    work(openPath, binPath, versionCnt, wordRangeBegin, wordRangeEnd, mode, questionCnt)
